"""
Resume parsing utilities to extract text from various file formats.
"""
from PyPDF2 import PdfReader
import docx
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def extract_text_from_pdf(path):
    """
    Extract text from PDF file.
    
    Args:
        path (str): Path to PDF file
        
    Returns:
        str: Extracted text
    """
    try:
        reader = PdfReader(path)
        texts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        
        extracted_text = "\n".join(texts)
        logger.info(f"Successfully extracted {len(extracted_text)} characters from PDF")
        return extracted_text
    
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(path):
    """
    Extract text from DOCX file.
    
    Args:
        path (str): Path to DOCX file
        
    Returns:
        str: Extracted text
    """
    try:
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        extracted_text = "\n".join(paragraphs)
        logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
        return extracted_text
    
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_txt(path):
    """
    Extract text from TXT file.
    
    Args:
        path (str): Path to TXT file
        
    Returns:
        str: Extracted text
    """
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        logger.info(f"Successfully extracted {len(text)} characters from TXT")
        return text
    
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        raise ValueError(f"Failed to extract text from TXT: {str(e)}")


def extract_text_from_file(path):
    """
    Extract text from file based on extension.
    Supports PDF, DOCX, DOC (as DOCX), and TXT files.
    
    Args:
        path (str): Path to file
        
    Returns:
        str: Extracted text
        
    Raises:
        ValueError: If file format is not supported or extraction fails
    """
    if not path:
        raise ValueError("File path cannot be empty")
    
    lower_path = path.lower()
    
    try:
        if lower_path.endswith('.pdf'):
            return extract_text_from_pdf(path)
        elif lower_path.endswith('.docx') or lower_path.endswith('.doc'):
            return extract_text_from_docx(path)
        elif lower_path.endswith('.txt'):
            return extract_text_from_txt(path)
        else:
            raise ValueError(f"Unsupported file format: {path.split('.')[-1]}")
    
    except Exception as e:
        logger.error(f"Error in extract_text_from_file: {e}")
        raise


def extract_resume_metadata(resume_text):
    """
    Extract metadata from resume text (email, phone, URLs, etc.).
    
    Args:
        resume_text (str): Resume text
        
    Returns:
        dict: Dictionary containing metadata
    """
    from utils import extract_email, extract_phone, extract_urls
    
    metadata = {
        'emails': extract_email(resume_text),
        'phones': extract_phone(resume_text),
        'urls': extract_urls(resume_text),
        'word_count': len(resume_text.split()),
        'char_count': len(resume_text)
    }
    
    return metadata
